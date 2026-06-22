"""Tests for :mod:`brailix.input.docx` — Word document adapter.

Fixtures are generated programmatically with ``python-docx`` + raw
lxml so we don't have to check binary ``.docx`` blobs into the repo;
each test builds the exact document shape it needs.

The whole module is skipped when ``python-docx`` isn't importable —
the adapter is gated on the ``docx`` extras group.
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest

pytest.importorskip("docx")
pytest.importorskip("lxml")

from docx import Document  # noqa: E402
from lxml import etree  # noqa: E402

from brailix.core import inline_math  # noqa: E402
from brailix.input.docx import parse_doc, parse_docx  # noqa: E402
from brailix.ir.document import (  # noqa: E402
    Heading,
    List,
    ListItem,
    MathBlock,
    Paragraph,
    Table,
)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_O_NS = "urn:schemas-microsoft-com:office:office"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


_INLINE_ISLAND_RE = re.compile(r"\$[^$\n]+\$")


def _inline_math_islands(text: str) -> list[str]:
    """Every inline-math island (a ``$...$`` region) in ``text`` — deferred
    source-tagged islands (inline OMML / EQ field) and eager
    ``$<math>...$`` ones (MTEF / script clusters) alike."""
    return _INLINE_ISLAND_RE.findall(text)


def _island_mathml(island: str) -> str:
    """MathML the frontend builds from a deferred tagged ``island``.

    Inline OMML / EQ math is no longer converted in the input layer; it
    travels as a source-tagged island that the frontend's math pass turns
    into a tree (exactly as ``Pipeline._attach_math`` does). Running that
    same conversion here lets a docx test assert the resulting MathML and
    so prove deferral is lossless.
    """
    import xml.etree.ElementTree as ET

    from brailix.core.context import MathContext
    from brailix.frontend.math import parse_math_tree

    source, payload = inline_math.unwrap(island)
    tree = parse_math_tree(payload, MathContext(source=source, profile="cn_current"))
    return "" if tree is None else ET.tostring(tree, encoding="unicode")


class TestResolveDocConverter:
    """``_resolve_doc_converter`` finds the LibreOffice binary (or None)."""

    def test_command_name_resolved_via_path(self, monkeypatch) -> None:
        # A bare command name (not a file in cwd) must be looked up on PATH,
        # not skipped — the operator-precedence bug returned None for it.
        from brailix.input.docx import _resolve_doc_converter

        monkeypatch.setattr(
            "brailix.input.docx.shutil.which",
            lambda name: "/usr/bin/soffice" if name == "soffice" else None,
        )
        assert _resolve_doc_converter("soffice") == "/usr/bin/soffice"

    def test_none_when_override_not_runnable(self, monkeypatch) -> None:
        from brailix.input.docx import _resolve_doc_converter

        monkeypatch.setattr(
            "brailix.input.docx.shutil.which", lambda name: None
        )
        assert _resolve_doc_converter("nope") is None


# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def _omml_fragment(xml_body: str) -> etree._Element:
    """Build an ``<m:oMath>`` element from a body XML snippet."""
    wrapped = (
        f'<m:oMath xmlns:m="{_M_NS}">{xml_body}</m:oMath>'
    )
    return etree.fromstring(wrapped)


def _omml_para(xml_body: str) -> etree._Element:
    """Build an ``<m:oMathPara>`` (display) element."""
    wrapped = (
        f'<m:oMathPara xmlns:m="{_M_NS}">'
        f'<m:oMath>{xml_body}</m:oMath>'
        f'</m:oMathPara>'
    )
    return etree.fromstring(wrapped)


def _make_docx(tmp_path: Path, name: str = "doc.docx") -> tuple[Path, Document]:
    """Create a fresh empty ``.docx`` at ``tmp_path / name``."""
    path = tmp_path / name
    doc = Document()
    return path, doc


# ---------------------------------------------------------------------------
# Plain paragraph / heading / list
# ---------------------------------------------------------------------------


class TestStructuralBlocks:
    def test_single_paragraph(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("我在重庆。")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        # python-docx adds a trailing blank paragraph by default
        # (it's the cursor position); filter empty paragraphs.
        non_empty = [
            b for b in result.blocks
            if not (isinstance(b, Paragraph) and not (b.text or "").strip())
        ]
        assert len(non_empty) == 1
        assert isinstance(non_empty[0], Paragraph)
        assert non_empty[0].text == "我在重庆。"

    def test_heading_style_becomes_heading_block(self, tmp_path: Path) -> None:
        # Headings are detected by the pStyle name. ``add_heading``
        # applies the "Heading 1" style which we recognise as level 1.
        path, doc = _make_docx(tmp_path)
        doc.add_heading("第一章", level=1)
        doc.add_heading("第一节", level=2)
        doc.add_paragraph("正文。")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        headings = [b for b in result.blocks if isinstance(b, Heading)]
        assert [h.level for h in headings] == [1, 2]
        assert [h.text for h in headings] == ["第一章", "第一节"]

    def test_list_paragraphs_group_into_one_list(self, tmp_path: Path) -> None:
        # python-docx's ``add_paragraph(style="List Bullet")`` attaches
        # the bullet-list style which triggers ``numPr``.
        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("一项", style="List Bullet")
        doc.add_paragraph("二项", style="List Bullet")
        doc.add_paragraph("三项", style="List Bullet")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        lists = [b for b in result.blocks if isinstance(b, List)]
        assert len(lists) == 1
        assert [it.text for it in lists[0].items] == ["一项", "二项", "三项"]
        assert all(isinstance(it, ListItem) for it in lists[0].items)

    def test_table_with_two_rows(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "甲"
        t.rows[0].cells[1].text = "乙"
        t.rows[1].cells[0].text = "丙"
        t.rows[1].cells[1].text = "丁"
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        tables = [b for b in result.blocks if isinstance(b, Table)]
        assert len(tables) == 1
        rows = tables[0].rows
        assert len(rows) == 2
        assert [c.text for c in rows[0].cells] == ["甲", "乙"]
        assert [c.text for c in rows[1].cells] == ["丙", "丁"]

    def test_nested_table_cell_content_preserved(self, tmp_path: Path) -> None:
        # A table nested inside a cell must not vanish — its inner cells'
        # text is folded into the parent cell text (old code skipped any
        # non-``p`` child, dropping the whole nested grid).
        path, doc = _make_docx(tmp_path)
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].add_run("外层")
        nested = cell.add_table(rows=1, cols=2)
        nested.cell(0, 0).text = "内甲"
        nested.cell(0, 1).text = "内乙"
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        tables = [b for b in result.blocks if isinstance(b, Table)]
        assert len(tables) == 1
        cell_text = tables[0].rows[0].cells[0].text or ""
        assert "外层" in cell_text
        assert "内甲" in cell_text
        assert "内乙" in cell_text


# ---------------------------------------------------------------------------
# List level (``w:numPr/w:ilvl``) parsing robustness
# ---------------------------------------------------------------------------


class TestNumPrLevel:
    """``w:numPr/w:ilvl`` level parsing tolerates malformed values instead
    of crashing — the old ``int(ilvl_elem.get(...))`` raised on a missing
    or unprefixed ``val``."""

    def _list_para(self, ilvl_xml: str):
        p_xml = (
            f'<w:p xmlns:w="{_W_NS}">'
            f'<w:pPr><w:numPr>{ilvl_xml}'
            f'<w:numId w:val="1"/></w:numPr></w:pPr>'
            f'<w:r><w:t>项</w:t></w:r>'
            f'</w:p>'
        )
        return etree.fromstring(p_xml)

    def test_missing_val_defaults_to_level_zero(self) -> None:
        # <w:ilvl/> with no val must not crash (old code did int(None)).
        from brailix.input.docx._blocks import _paragraph_list_info

        assert _paragraph_list_info(self._list_para("<w:ilvl/>")) == (0, False)

    def test_bare_val_without_prefix_is_read(self) -> None:
        # Some emitters write a bare ``val`` with no ``w:`` prefix.
        from brailix.input.docx._blocks import _paragraph_list_info

        assert _paragraph_list_info(
            self._list_para('<w:ilvl val="2"/>')
        ) == (2, False)

    def test_normal_namespaced_val(self) -> None:
        from brailix.input.docx._blocks import _paragraph_list_info

        assert _paragraph_list_info(
            self._list_para('<w:ilvl w:val="1"/>')
        ) == (1, False)

    def test_non_integer_val_defaults_to_zero(self) -> None:
        from brailix.input.docx._blocks import _paragraph_list_info

        assert _paragraph_list_info(
            self._list_para('<w:ilvl w:val="x"/>')
        ) == (0, False)


class TestParagraphStyle:
    """``w:pStyle@val`` reads the qualified form first, bare ``val`` as a
    fallback — some emitters drop the prefix, and reading only the qualified
    form silently lost heading / style-only list detection (the style name
    came back None, so the paragraph degraded to plain body text)."""

    def _styled_para(self, pstyle_xml: str):
        p_xml = (
            f'<w:p xmlns:w="{_W_NS}">'
            f"<w:pPr>{pstyle_xml}</w:pPr>"
            f"<w:r><w:t>标题</w:t></w:r>"
            f"</w:p>"
        )
        return etree.fromstring(p_xml)

    def test_normal_namespaced_val(self) -> None:
        from brailix.input.docx._blocks import _paragraph_style

        assert (
            _paragraph_style(self._styled_para('<w:pStyle w:val="Heading1"/>'))
            == "Heading1"
        )

    def test_bare_val_without_prefix_is_read(self) -> None:
        # Regression: a bare-val pStyle returned None, so heading / list-by-
        # style detection silently failed and the paragraph became body text.
        from brailix.input.docx._blocks import _paragraph_style

        assert (
            _paragraph_style(self._styled_para('<w:pStyle val="Heading1"/>'))
            == "Heading1"
        )


# ---------------------------------------------------------------------------
# Paragraph alignment (``w:jc`` → Block.align)
# ---------------------------------------------------------------------------


class TestParagraphAlignment:
    """A paragraph's ``w:jc`` survives as :attr:`Block.align`, but only for
    the alignments braille layout can honour (centre / right)."""

    def _only_paragraph(self, result) -> Paragraph:
        paras = [
            b for b in result.blocks
            if isinstance(b, Paragraph) and (b.text or "").strip()
        ]
        assert len(paras) == 1
        return paras[0]

    def test_centered_paragraph_carries_center(self, tmp_path: Path) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("居中标题")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(path)

        assert self._only_paragraph(parse_docx(path, profile="cn_current", language="zh-CN")).align == "center"

    def test_right_aligned_paragraph_carries_right(self, tmp_path: Path) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("二〇二六年五月")
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.save(path)

        assert self._only_paragraph(parse_docx(path, profile="cn_current", language="zh-CN")).align == "right"

    def test_justified_paragraph_has_no_align(self, tmp_path: Path) -> None:
        # Braille has no justification convention, so "both" normalises to
        # None — the paragraph reads flush-left, same as untagged prose.
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("一段两端对齐的正文内容")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.save(path)

        assert self._only_paragraph(parse_docx(path, profile="cn_current", language="zh-CN")).align is None

    def test_default_paragraph_has_no_align(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("普通左对齐段落")
        doc.save(path)

        assert self._only_paragraph(parse_docx(path, profile="cn_current", language="zh-CN")).align is None

    def test_centered_heading_carries_center(self, tmp_path: Path) -> None:
        # Alignment is recorded regardless of block kind: a centred level-2
        # heading carries align so the layout can centre it even though the
        # default rule centres only level 1.
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        path, doc = _make_docx(tmp_path)
        h = doc.add_heading("居中小标题", level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(path)

        headings = [b for b in parse_docx(path, profile="cn_current", language="zh-CN").blocks if isinstance(b, Heading)]
        assert len(headings) == 1
        assert headings[0].level == 2
        assert headings[0].align == "center"


class TestAlignmentEndToEnd:
    """parse → translate → layout: a centred Word paragraph renders centred."""

    def test_centered_paragraph_renders_centered(self, tmp_path: Path) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        from brailix.pipeline import Pipeline
        from brailix.renderer.layout import LayoutOptions, LayoutRenderer
        from brailix.renderer.unicode_braille import dots_to_char

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("一")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(path)

        ir = parse_docx(path, profile="cn_current", language="zh-CN")
        result = Pipeline(profile="cn_current").translate_document(ir)
        out = LayoutRenderer(options=LayoutOptions(line_width=20)).render(
            result.braille_ir
        )
        blank = dots_to_char(())
        content = [ln for ln in out.split("\n") if any(c != blank for c in ln)]
        assert content
        # Genuinely centred → leading padding well past the 2-cell first-line
        # indent a plain (flush-left) paragraph would have used.
        leading_blanks = len(content[0]) - len(content[0].lstrip(blank))
        assert leading_blanks > 2


# ---------------------------------------------------------------------------
# Math handling
# ---------------------------------------------------------------------------


class TestMath:
    def test_display_math_paragraph_becomes_math_block(
        self, tmp_path: Path
    ) -> None:
        # A paragraph that contains only ``<m:oMathPara>`` produces a
        # MathBlock; the OMML XML is preserved as ``MathBlock.text``.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        para._p.append(_omml_para(
            '<m:f><m:num><m:r><m:t>x</m:t></m:r></m:num>'
            '<m:den><m:r><m:t>2</m:t></m:r></m:den></m:f>'
        ))
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        math_blocks = [b for b in result.blocks if isinstance(b, MathBlock)]
        assert len(math_blocks) == 1
        assert math_blocks[0].source == "omml"
        # Round-trip preserved enough that the OMML adapter can still
        # find the fraction tags.
        assert "<m:f" in math_blocks[0].text or "m:f" in math_blocks[0].text

    def test_inline_math_embedded_in_paragraph_text(
        self, tmp_path: Path
    ) -> None:
        # Inline ``m:oMath`` (no ``m:oMathPara`` wrapper) stays in the
        # paragraph as a deferred source-tagged ``omml`` island, so the
        # frontend converts it later; the segmenter still sees it as ``$...$``.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("公式 ")
        para._p.append(_omml_fragment(
            '<m:sSup>'
            '<m:e><m:r><m:t>x</m:t></m:r></m:e>'
            '<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
            '</m:sSup>'
        ))
        # python-docx ``add_run`` after the math node appends text.
        para.add_run(" 是平方。")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        # The inline equation should sit in the same paragraph text as
        # the surrounding Chinese.
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "公式" in joined
        assert "是平方" in joined
        # Deferred: one tagged ``omml`` island carrying the raw OMML, which
        # the frontend converts to ``<msup>`` — not pre-converted here.
        islands = _inline_math_islands(joined)
        assert len(islands) == 1 and inline_math.is_tagged(islands[0])
        assert inline_math.unwrap(islands[0])[0] == "omml"
        assert "<msup>" in _island_mathml(islands[0])


class TestRevisionAndContentControlWrappers:
    """Revision-tracking (``<w:ins>`` / ``<w:del>`` …) and content-control
    (``<w:sdt>``) wrappers are transparent to content: the runs / inline math
    they wrap must parse exactly as if unwrapped, not be scraped to bare text.
    Tracked changes is a routine real-document state (review / accessibility
    remediation workflows), so a formula inside an insertion must not be
    silently flattened to literal characters."""

    def test_inline_omml_inside_tracked_insertion_survives(
        self, tmp_path: Path
    ) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("公式 ")
        ins = etree.fromstring(
            f'<w:ins xmlns:w="{_W_NS}" xmlns:m="{_M_NS}" '
            f'w:id="7" w:author="ed" w:date="2024-01-01T00:00:00Z">'
            f'<m:oMath><m:sSup>'
            f'<m:e><m:r><m:t>x</m:t></m:r></m:e>'
            f'<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
            f'</m:sSup></m:oMath>'
            f'</w:ins>'
        )
        para._p.append(ins)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        joined = "\n".join(
            p.text or "" for p in result.blocks if isinstance(p, Paragraph)
        )
        islands = _inline_math_islands(joined)
        assert len(islands) == 1 and inline_math.is_tagged(islands[0])
        assert inline_math.unwrap(islands[0])[0] == "omml"
        assert "<msup>" in _island_mathml(islands[0])

    def test_inline_omml_inside_content_control_survives(
        self, tmp_path: Path
    ) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("公式 ")
        sdt = etree.fromstring(
            f'<w:sdt xmlns:w="{_W_NS}" xmlns:m="{_M_NS}"><w:sdtContent>'
            f'<m:oMath><m:sSup>'
            f'<m:e><m:r><m:t>x</m:t></m:r></m:e>'
            f'<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
            f'</m:sSup></m:oMath>'
            f'</w:sdtContent></w:sdt>'
        )
        para._p.append(sdt)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        joined = "\n".join(
            p.text or "" for p in result.blocks if isinstance(p, Paragraph)
        )
        islands = _inline_math_islands(joined)
        assert len(islands) == 1 and inline_math.is_tagged(islands[0])
        assert "<msup>" in _island_mathml(islands[0])

    def test_inline_math_under_unknown_wrapper_survives(
        self, tmp_path: Path
    ) -> None:
        # docx-unknown-wrapper: inline math under an unrecognised wrapping
        # element (not in _TRANSPARENT_RUN_WRAPPERS) survives via transparent
        # descent, not flattened to its bare <w:t> characters.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("公式 ")
        wrap = etree.fromstring(
            f'<w:futureWrapper xmlns:w="{_W_NS}" xmlns:m="{_M_NS}">'
            f"<m:oMath><m:sSup>"
            f"<m:e><m:r><m:t>x</m:t></m:r></m:e>"
            f"<m:sup><m:r><m:t>2</m:t></m:r></m:sup>"
            f"</m:sSup></m:oMath>"
            f"</w:futureWrapper>"
        )
        para._p.append(wrap)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        joined = "\n".join(
            p.text or "" for p in result.blocks if isinstance(p, Paragraph)
        )
        islands = _inline_math_islands(joined)
        assert len(islands) == 1
        assert "<msup>" in _island_mathml(islands[0])

    def test_run_nested_alternate_content_display_math_survives(
        self, tmp_path: Path
    ) -> None:
        # docx-altcontent: a display equation (oMathPara) in an
        # AlternateContent nested inside a <w:r> was dropped (its _math slot
        # discarded). It now folds in as an inline-math island, not vanishes.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("公式 ")
        run = etree.fromstring(
            f'<w:r xmlns:w="{_W_NS}" xmlns:m="{_M_NS}" xmlns:mc="{_MC_NS}">'
            f"<mc:AlternateContent><mc:Fallback>"
            f"<m:oMathPara><m:oMath><m:sSup>"
            f"<m:e><m:r><m:t>x</m:t></m:r></m:e>"
            f"<m:sup><m:r><m:t>2</m:t></m:r></m:sup>"
            f"</m:sSup></m:oMath></m:oMathPara>"
            f"</mc:Fallback></mc:AlternateContent>"
            f"</w:r>"
        )
        para._p.append(run)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        joined = "\n".join(
            p.text or "" for p in result.blocks if isinstance(p, Paragraph)
        )
        islands = _inline_math_islands(joined)
        assert len(islands) == 1
        assert "<msup>" in _island_mathml(islands[0])

    def test_block_paragraph_inside_content_control_not_dropped(
        self, tmp_path: Path
    ) -> None:
        # A whole paragraph wrapped in a block-level content control must not
        # vanish (the old ``else: continue`` dropped every non-p/tbl child).
        path, doc = _make_docx(tmp_path)
        sdt = etree.fromstring(
            f'<w:sdt xmlns:w="{_W_NS}"><w:sdtContent>'
            f'<w:p><w:r><w:t>内容控件中的整段文字。</w:t></w:r></w:p>'
            f'</w:sdtContent></w:sdt>'
        )
        doc.element.body.insert(0, sdt)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        joined = "\n".join(
            p.text or "" for p in result.blocks if isinstance(p, Paragraph)
        )
        assert "内容控件中的整段文字" in joined

    def test_table_row_inside_tracked_insertion_not_dropped(
        self, tmp_path: Path
    ) -> None:
        # A whole table row wrapped in <w:ins> (an inserted row, common in
        # accept-changes-first documents) must not be silently dropped — the
        # body walker descends these wrappers, the table walker now does too.
        path, doc = _make_docx(tmp_path)
        tbl = etree.fromstring(
            f'<w:tbl xmlns:w="{_W_NS}">'
            f"<w:tr><w:tc><w:p><w:r><w:t>裸行</w:t></w:r></w:p></w:tc></w:tr>"
            f'<w:ins w:id="9" w:author="ed" w:date="2024-01-01T00:00:00Z">'
            f"<w:tr><w:tc><w:p><w:r><w:t>插入行</w:t></w:r></w:p></w:tc></w:tr>"
            f"</w:ins>"
            f"</w:tbl>"
        )
        doc.element.body.insert(0, tbl)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        tables = [b for b in result.blocks if isinstance(b, Table)]
        assert tables
        cell_text = "\n".join(
            c.text or "" for row in tables[0].rows for c in row.cells
        )
        assert "裸行" in cell_text
        assert "插入行" in cell_text

    def test_table_cell_inside_tracked_insertion_not_dropped(
        self, tmp_path: Path
    ) -> None:
        # A single cell wrapped in <w:ins> within a row must survive too.
        path, doc = _make_docx(tmp_path)
        tbl = etree.fromstring(
            f'<w:tbl xmlns:w="{_W_NS}">'
            f"<w:tr>"
            f"<w:tc><w:p><w:r><w:t>甲</w:t></w:r></w:p></w:tc>"
            f'<w:ins w:id="10" w:author="ed" w:date="2024-01-01T00:00:00Z">'
            f"<w:tc><w:p><w:r><w:t>乙</w:t></w:r></w:p></w:tc>"
            f"</w:ins>"
            f"</w:tr>"
            f"</w:tbl>"
        )
        doc.element.body.insert(0, tbl)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        tables = [b for b in result.blocks if isinstance(b, Table)]
        assert tables
        cell_text = "\n".join(
            c.text or "" for row in tables[0].rows for c in row.cells
        )
        assert "甲" in cell_text and "乙" in cell_text


# ---------------------------------------------------------------------------
# MathType / Equation 3.0 OLE handling
# ---------------------------------------------------------------------------


def _embed_ole_equation(
    doc: Document, paragraph, blob: bytes, *, progid: str = "Equation.DSMT4"
) -> str:
    """Attach ``blob`` as an OLE equation part and reference it from ``paragraph``.

    Inserts a minimal ``<w:object>`` + ``<o:OLEObject>`` skeleton inside
    a new ``<w:r>`` at the end of ``paragraph``, and creates the matching
    ``OLE_OBJECT`` relationship so ``parse_docx`` can resolve the rId
    back to the part blob. Returns the assigned rId for visibility.

    The OLE part's content type and partname mirror what Word emits;
    python-docx's ``Part`` constructor + ``relate_to`` does the rest of
    the bookkeeping (rels XML, [Content_Types].xml, package zip layout).
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    # python-docx assigns unique rIds across parts; pick the next
    # available bin number too so multiple equations don't collide.
    existing = [
        n for n in doc.part.package.iter_parts()
        if "/word/embeddings/oleObject" in n.partname
    ]
    idx = len(existing) + 1
    ole_part = Part(
        partname=PackURI(f"/word/embeddings/oleObject{idx}.bin"),
        content_type="application/vnd.openxmlformats-officedocument.oleObject",
        blob=blob,
        package=doc.part.package,
    )
    rid = doc.part.relate_to(ole_part, RT.OLE_OBJECT)

    # Build the in-paragraph markup. Word's typical layout is:
    #   <w:r>
    #     <w:object>
    #       <o:OLEObject Type="Embed" ProgID="..." r:id="...".../>
    #     </w:object>
    #   </w:r>
    obj_xml = (
        f'<w:r xmlns:w="{_W_NS}" xmlns:o="{_O_NS}" xmlns:r="{_R_NS}">'
        f'<w:object>'
        f'<o:OLEObject Type="Embed" ProgID="{progid}" '
        f'ShapeID="1000" DrawAspect="Content" ObjectID="1" '
        f'r:id="{rid}"/>'
        f'</w:object>'
        f'</w:r>'
    )
    paragraph._p.append(etree.fromstring(obj_xml))
    return rid


def _relate_ole_blob(
    doc: Document, blob: bytes, *, progid: str = "Equation.DSMT4"
) -> str:
    """Create an OLE part for ``blob`` + an OLE_OBJECT relationship; return its rId.

    Same part / relationship bookkeeping as :func:`_embed_ole_equation` but
    *without* inserting the in-paragraph ``<w:object>`` markup — the caller
    places the reference itself (e.g. inside an ``<mc:AlternateContent>``
    Fallback branch). ``progid`` is unused for the part itself (the ProgID
    lives on the caller's ``<o:OLEObject>``) but kept for call-site symmetry.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    existing = [
        n for n in doc.part.package.iter_parts()
        if "/word/embeddings/oleObject" in n.partname
    ]
    idx = len(existing) + 1
    ole_part = Part(
        partname=PackURI(f"/word/embeddings/oleObject{idx}.bin"),
        content_type="application/vnd.openxmlformats-officedocument.oleObject",
        blob=blob,
        package=doc.part.package,
    )
    return doc.part.relate_to(ole_part, RT.OLE_OBJECT)


def _mtef_sample_payload() -> bytes:
    """Hand-crafted MTEF v5 payload encoding ``x²``.

    Reuses the same byte-builder the math-adapter tests use so any
    selector regression shows up here AND in the adapter tests.
    """
    from tests.frontend.math._mtef_builder import (
        v5_line,
        v5_prelude,
        v5_simple_char_line,
        v5_tmpl,
    )
    base = v5_simple_char_line(ord("x"))
    sup = v5_simple_char_line(ord("2"))
    # selector 28 = superscript
    return v5_prelude() + v5_line(v5_tmpl(28, [base, sup]))


def _eqnolehdr_wrapped(payload: bytes) -> bytes:
    """Prepend a 28-byte EQNOLEFILEHDR so the blob mimics an OLE stream."""
    header = (
        bytes([0x1C, 0x00])  # cbHdr=28
        + bytes([0x00, 0x00, 0x02, 0x00])  # version
        + bytes([0x00, 0x00])  # cf
        + len(payload).to_bytes(4, "little")  # cbObject
        + bytes(16)  # reserved
    )
    assert len(header) == 28
    return header + payload


class TestMathTypeOLE:
    def test_inline_ole_equation_becomes_inline_math(
        self, tmp_path: Path
    ) -> None:
        # A paragraph containing "公式 [OLE x²] 是平方" — the docx adapter
        # should resolve the OLE rId, run MTEF→MathML, and surface the
        # result as inline $<math>...</math>$ text in the paragraph.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("公式 ")
        _embed_ole_equation(
            doc, para, _eqnolehdr_wrapped(_mtef_sample_payload())
        )
        para.add_run(" 是平方。")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "公式" in joined
        assert "是平方" in joined
        assert "$<math" in joined and "</math>$" in joined
        assert "<msup>" in joined
        assert "<mi>x</mi>" in joined
        assert "<mn>2</mn>" in joined

    def test_inline_ole_equation_real_mathtype_bytes(
        self, tmp_path: Path
    ) -> None:
        # End-to-end with REAL MathType bytes (the committed Equation
        # Native fixture), not the synthetic slot-0 builder: docx OLE →
        # MTEF decode → MathML must reproduce y = x³ with the base x
        # *inside* the <msup> (the preceding-sibling fix), proving the
        # docx path handles the real wire shape — the synthetic builder
        # puts the base in slot 0, which masked this exact bug class once.
        import pathlib

        blob = (
            pathlib.Path(__file__).parent.parent
            / "frontend" / "math" / "fixtures"
            / "mathtype_v5_y_eq_x_cubed.bin"
        ).read_bytes()  # already a full Equation Native stream (EQNOLEFILEHDR)
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("结果 ")
        _embed_ole_equation(doc, para, blob)
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        joined = "\n".join(
            p.text or "" for p in result.blocks if isinstance(p, Paragraph)
        )
        assert "$<math" in joined and "</math>$" in joined
        assert "<mi>y</mi>" in joined
        assert "<mo>=</mo>" in joined
        # Base x lives inside the script, not orphaned as a sibling.
        assert "<msup><mi>x</mi><mn>3</mn></msup>" in joined

    def test_equation_3_progid_also_recognised(
        self, tmp_path: Path
    ) -> None:
        # ProgID "Equation.3" is the legacy Microsoft Equation 3.0
        # marker; many old textbooks use it. Same MTEF wire format.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("旧公式 ")
        _embed_ole_equation(
            doc, para,
            _eqnolehdr_wrapped(_mtef_sample_payload()),
            progid="Equation.3",
        )
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "<msup>" in joined

    def test_unknown_progid_is_skipped(self, tmp_path: Path) -> None:
        # An OLE object that isn't an equation (e.g. an embedded
        # spreadsheet) must not crash parse_docx — we just leave it
        # out of the IR.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("文本 ")
        _embed_ole_equation(
            doc, para, b"random bytes", progid="Excel.Sheet.12"
        )
        para.add_run(" 末尾")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "文本" in joined
        assert "末尾" in joined
        assert "$<math" not in joined

    def test_malformed_mtef_yields_merror_not_crash(
        self, tmp_path: Path
    ) -> None:
        # Truncated MTEF payload still has to produce a recognisable
        # inline math marker (with merror inside) — the document must
        # round-trip end-to-end without throwing.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("坏 ")
        # First byte is a v5 version, but the rest is junk → adapter
        # raises internally and falls back to merror_wrap.
        _embed_ole_equation(doc, para, bytes([5, 1, 0, 11, 0, 0xFF, 0xFF]))
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "$<math" in joined
        assert "merror" in joined

    def test_oversized_ole_blob_is_skipped(self, monkeypatch) -> None:
        # docx-mtef: an oversized untrusted embed is skipped (treated as
        # non-math) rather than handed to the OLE / MTEF parsers, so a hostile
        # or corrupt .docx can't inflate memory through one stream.
        from brailix.input.docx import _ole

        monkeypatch.setattr(_ole, "_MAX_MTEF_BYTES", 16)
        # Over the cap → skipped, even though the prelude looks like raw MTEF.
        assert _ole._extract_mtef_payload(b"\x05" + b"\x00" * 32) is None
        # Within the cap → the raw-MTEF heuristic still recognises it.
        small = b"\x05" + b"\x00" * 8
        assert _ole._extract_mtef_payload(small) == small

    def test_no_ole_objects_path_unaffected(self, tmp_path: Path) -> None:
        # Sanity: a document without any OLE objects produces no
        # ole_blobs entries and should walk identically to the pre-
        # patch behaviour.
        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("纯文本")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        # Build the blob map manually to confirm it's empty.
        from brailix.input.docx import _build_ole_blob_map
        assert _build_ole_blob_map(doc) == {}
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        non_empty = [p for p in paragraphs if (p.text or "").strip()]
        assert len(non_empty) == 1
        assert non_empty[0].text == "纯文本"

    def test_external_ole_relationship_does_not_crash(
        self, tmp_path: Path
    ) -> None:
        # A *linked* (not embedded) OLE object is an external relationship.
        # python-docx raises ValueError — NOT AttributeError — when such a
        # relationship is asked for its ``target_part``; the blob-map builder
        # must skip it rather than let that escape and crash the whole parse.
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        from brailix.input.docx import _build_ole_blob_map

        _, doc = _make_docx(tmp_path)
        doc.add_paragraph("链接对象")
        rid = doc.part.relate_to(
            "file:///C:/linked/equation.bin", RT.OLE_OBJECT, is_external=True
        )
        # Must not raise; the external rel contributes no blob (no local part).
        blob_map = _build_ole_blob_map(doc)
        assert rid not in blob_map


# ---------------------------------------------------------------------------
# mc:AlternateContent (Word offers a modern Choice + legacy Fallback)
# ---------------------------------------------------------------------------


class TestAlternateContent:
    """Word wraps an OLE object in ``<mc:AlternateContent>`` when it can also
    offer a modern (Choice) representation. The adapter prefers the Fallback's
    legacy ``<w:object>`` OLE equation — that's the one it can read — and only
    descends into Choice when Fallback yields nothing. The whole recursive
    ``_walk_alternate_content`` / ``_walk_alt_subtree`` path was previously
    untested."""

    def test_fallback_ole_equation_is_extracted(self, tmp_path: Path) -> None:
        # Fallback holds the MathType OLE; Choice holds a preview we can't read.
        # The OLE must surface as inline math and the Choice preview must NOT
        # leak into the text (Fallback is preferred and short-circuits Choice).
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("前 ")
        rid = _relate_ole_blob(doc, _eqnolehdr_wrapped(_mtef_sample_payload()))
        alt_xml = (
            f'<mc:AlternateContent xmlns:mc="{_MC_NS}" xmlns:w="{_W_NS}" '
            f'xmlns:o="{_O_NS}" xmlns:r="{_R_NS}">'
            f'<mc:Choice Requires="wps"><w:r><w:t>预览图</w:t></w:r></mc:Choice>'
            f'<mc:Fallback><w:r><w:object>'
            f'<o:OLEObject Type="Embed" ProgID="Equation.DSMT4" r:id="{rid}"/>'
            f'</w:object></w:r></mc:Fallback>'
            f'</mc:AlternateContent>'
        )
        para._p.append(etree.fromstring(alt_xml))
        doc.save(path)

        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "前" in text
        assert "$<math" in text and "</math>$" in text
        assert "<msup>" in text
        assert "<mi>x</mi>" in text and "<mn>2</mn>" in text
        # Choice preview is dropped — Fallback won, so we never walked it.
        assert "预览图" not in text

    def test_fallback_object_direct_child_also_extracted(
        self, tmp_path: Path
    ) -> None:
        # Some emitters put the ``<w:object>`` directly under Fallback with no
        # wrapping ``<w:r>`` — the subtree walker must still find it.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        rid = _relate_ole_blob(doc, _eqnolehdr_wrapped(_mtef_sample_payload()))
        alt_xml = (
            f'<mc:AlternateContent xmlns:mc="{_MC_NS}" xmlns:w="{_W_NS}" '
            f'xmlns:o="{_O_NS}" xmlns:r="{_R_NS}">'
            f'<mc:Fallback><w:object>'
            f'<o:OLEObject Type="Embed" ProgID="Equation.3" r:id="{rid}"/>'
            f'</w:object></mc:Fallback>'
            f'</mc:AlternateContent>'
        )
        para._p.append(etree.fromstring(alt_xml))
        doc.save(path)

        assert "<msup>" in _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))

    def test_choice_only_omml_is_extracted(self, tmp_path: Path) -> None:
        # No Fallback at all → the adapter descends into Choice and surfaces
        # its inline OMML as math (the Choice branch of _walk_alternate_content).
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("式 ")
        alt_xml = (
            f'<mc:AlternateContent xmlns:mc="{_MC_NS}" xmlns:w="{_W_NS}" '
            f'xmlns:m="{_M_NS}">'
            f'<mc:Choice Requires="wps">'
            f'<m:oMath><m:sSup>'
            f'<m:e><m:r><m:t>x</m:t></m:r></m:e>'
            f'<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
            f'</m:sSup></m:oMath>'
            f'</mc:Choice>'
            f'</mc:AlternateContent>'
        )
        para._p.append(etree.fromstring(alt_xml))
        doc.save(path)

        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "式" in text
        # The Choice's inline OMML is deferred, like any other inline OMML.
        islands = _inline_math_islands(text)
        assert len(islands) == 1 and inline_math.unwrap(islands[0])[0] == "omml"
        assert "<msup>" in _island_mathml(islands[0])

    def test_empty_fallback_does_not_suppress_choice_math(self, tmp_path: Path) -> None:
        # Fallback holds only an empty placeholder run; the formula lives in
        # Choice. The empty Fallback must NOT short-circuit Choice — regression:
        # an empty ("", None) piece used to set produced=True and drop the
        # Choice formula.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("式 ")
        alt_xml = (
            f'<mc:AlternateContent xmlns:mc="{_MC_NS}" xmlns:w="{_W_NS}" '
            f'xmlns:m="{_M_NS}">'
            f'<mc:Choice Requires="wps">'
            f'<m:oMath><m:sSup>'
            f'<m:e><m:r><m:t>x</m:t></m:r></m:e>'
            f'<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
            f'</m:sSup></m:oMath>'
            f'</mc:Choice>'
            f'<mc:Fallback><w:r><w:t></w:t></w:r></mc:Fallback>'
            f'</mc:AlternateContent>'
        )
        para._p.append(etree.fromstring(alt_xml))
        doc.save(path)
        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        islands = _inline_math_islands(text)
        assert len(islands) == 1  # Choice formula survived the empty Fallback
        assert "<msup>" in _island_mathml(islands[0])

    def test_cross_run_eq_field_in_alternate_content_assembles(
        self, tmp_path: Path
    ) -> None:
        # A cross-run EQ field (fldChar begin / instrText / end over 3 runs)
        # inside an AlternateContent branch must assemble into a formula, not be
        # dropped: _walk_alt_subtree now threads a _FieldState through its runs
        # (previously fldChar was skipped with no state). Regression.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("公式 ")
        eq_runs = (
            f'<w:r xmlns:w="{_W_NS}"><w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r xmlns:w="{_W_NS}"><w:instrText xml:space="preserve">'
            f" EQ \\F(1,2) </w:instrText></w:r>"
            f'<w:r xmlns:w="{_W_NS}"><w:fldChar w:fldCharType="end"/></w:r>'
        )
        alt_xml = (
            f'<mc:AlternateContent xmlns:mc="{_MC_NS}" xmlns:w="{_W_NS}">'
            f'<mc:Choice Requires="wps">{eq_runs}</mc:Choice>'
            f"</mc:AlternateContent>"
        )
        para._p.append(etree.fromstring(alt_xml))
        doc.save(path)
        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        islands = _inline_math_islands(text)
        assert len(islands) == 1  # the EQ field assembled, not dropped
        assert "<mfrac" in _island_mathml(islands[0])

    def test_recursion_error_in_body_walk_becomes_parse_error(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # A pathologically nested docx blows the Python stack in the block
        # walkers; the RecursionError must surface as ParseError (the
        # malformed-docx contract), not escape raw. Simulated at the walk seam.
        import brailix.input.docx as docx_mod
        from brailix.core.errors import ParseError

        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("x")
        doc.save(path)

        def _boom(*_a, **_k):
            raise RecursionError("simulated pathological nesting")

        monkeypatch.setattr(docx_mod, "_iter_body_blocks", _boom)
        with pytest.raises(ParseError):
            parse_docx(path, profile="cn_current", language="zh-CN")


# ---------------------------------------------------------------------------
# Word EQ field (legacy ``eq \\f(...)`` style equations)
# ---------------------------------------------------------------------------


def _embed_eq_field(paragraph, instr: str) -> None:
    """Append a Word EQ field to ``paragraph`` using the multi-run form.

    Builds the canonical ``fldChar(begin) → instrText → fldChar(end)``
    sequence inside three runs, with an optional ``separate`` + cached
    result that the parser must ignore.
    """
    runs_xml = (
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:instrText xml:space="preserve">{instr}</w:instrText>'
        f'</w:r>'
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )
    # ``etree.fromstring`` accepts only one root, so wrap and unpack.
    wrapper = etree.fromstring(f'<root xmlns:w="{_W_NS}">{runs_xml}</root>')
    for run in list(wrapper):
        paragraph._p.append(run)


def _embed_eq_field_with_result(paragraph, instr: str, cached: str) -> None:
    """EQ field with a ``separate`` + cached-result section."""
    runs_xml = (
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:instrText xml:space="preserve">{instr}</w:instrText>'
        f'</w:r>'
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:t>{cached}</w:t>'
        f'</w:r>'
        f'<w:r xmlns:w="{_W_NS}">'
        f'<w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )
    wrapper = etree.fromstring(f'<root xmlns:w="{_W_NS}">{runs_xml}</root>')
    for run in list(wrapper):
        paragraph._p.append(run)


class TestEqField:
    def test_simple_fraction_becomes_inline_math(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("分数 ")
        _embed_eq_field(para, r"eq \f(1,2)")
        para.add_run(" 是二分之一")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "分数" in joined
        assert "是二分之一" in joined
        # Deferred: one tagged ``eq_field`` island; the frontend converts it.
        islands = _inline_math_islands(joined)
        assert len(islands) == 1 and inline_math.unwrap(islands[0])[0] == "eq_field"
        mathml = _island_mathml(islands[0])
        assert "<mfrac>" in mathml
        assert "<mn>1</mn>" in mathml
        assert "<mn>2</mn>" in mathml

    def test_unclosed_field_does_not_eat_rest_of_paragraph(
        self, tmp_path: Path
    ) -> None:
        # A field with begin+separate but NO end (truncated / corrupt / split
        # by a revision) used to silently drop everything after `separate`.
        # The visible-fallback text must be recovered, not eaten.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("BEFORE")
        runs_xml = (
            f'<w:r xmlns:w="{_W_NS}"><w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r xmlns:w="{_W_NS}"><w:instrText xml:space="preserve">'
            f"eq \\f(1,2)</w:instrText></w:r>"
            f'<w:r xmlns:w="{_W_NS}"><w:fldChar w:fldCharType="separate"/></w:r>'
            f'<w:r xmlns:w="{_W_NS}"><w:t>RESULT</w:t></w:r>'
        )
        wrapper = etree.fromstring(f'<root xmlns:w="{_W_NS}">{runs_xml}</root>')
        for run in list(wrapper):
            para._p.append(run)
        para.add_run("AFTER_TAIL")
        doc.save(path)

        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "BEFORE" in text
        assert "AFTER_TAIL" in text  # no longer swallowed by the open field

    def test_cross_paragraph_field_keeps_visible_text(
        self, tmp_path: Path
    ) -> None:
        # A field whose begin/instr are in one paragraph and `end` in the next
        # (complex / revision-split fields) used to drop the first paragraph's
        # visible-fallback text entirely. It must be recovered as text.
        path, doc = _make_docx(tmp_path)
        p1 = doc.add_paragraph("P1HEAD")
        runs_xml = (
            f'<w:r xmlns:w="{_W_NS}"><w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r xmlns:w="{_W_NS}"><w:instrText xml:space="preserve">'
            f"eq \\f(1,2)</w:instrText></w:r>"
            f'<w:r xmlns:w="{_W_NS}"><w:fldChar w:fldCharType="separate"/></w:r>'
            f'<w:r xmlns:w="{_W_NS}"><w:t>P1TAIL</w:t></w:r>'
        )
        wrapper = etree.fromstring(f'<root xmlns:w="{_W_NS}">{runs_xml}</root>')
        for run in list(wrapper):
            p1._p.append(run)
        p2 = doc.add_paragraph("P2HEAD")
        p2._p.append(
            etree.fromstring(
                f'<w:r xmlns:w="{_W_NS}"><w:fldChar w:fldCharType="end"/></w:r>'
            )
        )
        doc.save(path)

        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "P1HEAD" in text
        assert "P1TAIL" in text  # first paragraph's visible text recovered
        assert "P2HEAD" in text

    def test_piecewise_function_from_problem_15(self, tmp_path: Path) -> None:
        # The actual EQ field text from ``周练习6-5.4学生版.docx`` problem
        # 15 — the original regression that motivated this whole path.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("f(x)＝")
        _embed_eq_field(
            para,
            r"eq \b\lc\{(\a\vs4\al\co1(sin x，x≥0，,x＋2，x&lt;0，))",
        )
        para.add_run("则不等式 f(x)>")
        _embed_eq_field(para, r"eq \f(1,2)")
        para.add_run(" 的解集是____.")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        # Surrounding Chinese text survived.
        assert "f(x)＝" in joined
        assert "解集" in joined
        # Two deferred eq_field islands — the piecewise function and the 1/2.
        islands = _inline_math_islands(joined)
        assert len(islands) == 2
        assert all(inline_math.unwrap(i)[0] == "eq_field" for i in islands)
        mathml = "".join(_island_mathml(i) for i in islands)
        # Piecewise: left brace, no right brace, mtable with 2 rows.
        assert '<mo fence="true">{</mo>' in mathml
        assert '<mo fence="true">}</mo>' not in mathml
        assert "<mtable" in mathml
        # The fraction.
        assert "<mfrac>" in mathml

    def test_cached_result_is_dropped(self, tmp_path: Path) -> None:
        # When a field has a ``separate`` + cached result, the result
        # text must NOT show up alongside the converted math — Word
        # writes a visual fallback like "x²" there, which would be
        # mistaken for paragraph text if we kept it.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _embed_eq_field_with_result(
            para, r"eq \f(1,2)", "CACHED_RESULT_TEXT"
        )
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "CACHED_RESULT_TEXT" not in joined
        islands = _inline_math_islands(joined)
        assert len(islands) == 1
        assert "<mfrac>" in _island_mathml(islands[0])

    def test_non_eq_field_is_skipped_silently(self, tmp_path: Path) -> None:
        # HYPERLINK / PAGE / TOC and the rest of Word's non-equation
        # fields must not throw, must not produce math, and must let
        # surrounding text through unchanged.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("前")
        _embed_eq_field_with_result(
            para, r"HYPERLINK \"http://example.com\"", "click here"
        )
        para.add_run("后")
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "前" in joined
        assert "后" in joined
        assert "$<math" not in joined
        # The cached result is dropped (we don't try to display
        # hyperlink text — that's a separate feature).
        assert "click here" not in joined

    def test_fld_simple_form_also_supported(self, tmp_path: Path) -> None:
        # Word's ``<w:fldSimple w:instr="eq \\f(1,2)">`` shortcut form
        # is rarer but still used by some emitters.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("简 ")
        fld_xml = (
            f'<w:fldSimple xmlns:w="{_W_NS}" w:instr="eq \\f(1,2)">'
            f'<w:r><w:t>cached</w:t></w:r>'
            f'</w:fldSimple>'
        )
        para._p.append(etree.fromstring(fld_xml))
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        islands = _inline_math_islands(joined)
        assert len(islands) == 1 and inline_math.unwrap(islands[0])[0] == "eq_field"
        assert "<mfrac>" in _island_mathml(islands[0])


# ---------------------------------------------------------------------------
# mathtype_fallback — LibreOffice escape hatch
# ---------------------------------------------------------------------------


class TestMathTypeFallback:
    def test_invalid_fallback_value_rejected(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("x")
        doc.save(path)
        with pytest.raises(ValueError):
            parse_docx(path, mathtype_fallback="bogus", profile="cn_current", language="zh-CN")

    def test_libreoffice_mode_without_converter_raises(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        from brailix.core.errors import ParseError

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: None,
        )
        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("x")
        doc.save(path)
        with pytest.raises(ParseError) as exc:
            parse_docx(path, mathtype_fallback="libreoffice", profile="cn_current", language="zh-CN")
        assert "LibreOffice" in str(exc.value)

    def test_libreoffice_mode_invokes_converter(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # When mathtype_fallback="libreoffice" is set, we should
        # delegate the entire parse to the LibreOffice path even if
        # the document has no OLE objects.
        calls: list[list[str]] = []

        def fake_run(cmd, *, check, capture_output, timeout):
            calls.append(cmd)
            # Simulate LibreOffice writing the converted file.
            out_dir = Path(cmd[cmd.index("--outdir") + 1])
            stem = Path(cmd[-1]).stem
            (out_dir / f"{stem}.docx").write_bytes(
                Path(cmd[-1]).read_bytes()
            )

            class Result:
                returncode = 0
            return Result()

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: "soffice",
        )
        monkeypatch.setattr("brailix.input.docx.subprocess.run", fake_run)

        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("回归测试")
        doc.save(path)

        result = parse_docx(path, mathtype_fallback="libreoffice", profile="cn_current", language="zh-CN")
        assert len(calls) == 1
        assert "--convert-to" in calls[0]
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "回归测试" in joined

    def test_auto_mode_skips_libreoffice_when_native_succeeds(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # If the native MTEF adapter parses everything cleanly,
        # mathtype_fallback="auto" must NOT call LibreOffice.
        calls: list[list[str]] = []

        def fake_run(*args, **kwargs):
            calls.append(args)
            raise AssertionError("LibreOffice should not be invoked")

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: "soffice",
        )
        monkeypatch.setattr("brailix.input.docx.subprocess.run", fake_run)

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("前 ")
        _embed_ole_equation(
            doc, para, _eqnolehdr_wrapped(_mtef_sample_payload())
        )
        doc.save(path)

        result = parse_docx(path, mathtype_fallback="auto", profile="cn_current", language="zh-CN")
        assert calls == []
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "<msup>" in joined

    def test_recovery_needed_unit_semantics(self) -> None:
        # Unit semantics of the auto-retry decision: silent loss (fewer
        # spans than equation OLEs) and all-merror both trigger; a
        # healthy span count with at least one success doesn't.
        from brailix.input.docx import _mtef_recovery_needed
        from brailix.ir.document import DocumentIR

        def doc_with(text: str) -> DocumentIR:
            return DocumentIR(blocks=[Paragraph(text=text)])

        ok = "$<math><mi>x</mi></math>$"
        bad = "$<math><merror><mtext>?</mtext></merror></math>$"

        # One equation OLE, zero spans — it vanished silently → retry.
        assert _mtef_recovery_needed(doc_with("no math here"), 1) is True
        # Two equation OLEs, only one span — one vanished → retry.
        assert _mtef_recovery_needed(doc_with(f"前 {ok}"), 2) is True
        # Every span is a soft failure → retry.
        assert _mtef_recovery_needed(doc_with(f"前 {bad}"), 1) is True
        # As many spans as equations, at least one decoded → no retry.
        assert _mtef_recovery_needed(doc_with(f"前 {ok} 后 {bad}"), 2) is False

    def test_recovery_counts_spans_nested_in_table_and_list(self) -> None:
        # An equation living in a table cell (or list item) must count
        # toward the span total exactly like a top-level one. The span
        # lives on the *child* block's text, so a flat result.blocks walk
        # missed it — making "auto" read the native decode as a silent loss
        # and do a needless LibreOffice round-trip.
        from brailix.input.docx import _mtef_recovery_needed
        from brailix.ir.document import DocumentIR, TableCell, TableRow

        ok = "$<math><mi>x</mi></math>$"
        in_table = DocumentIR(
            blocks=[Table(rows=[TableRow(cells=[TableCell(text=f"前 {ok}")])])]
        )
        in_list = DocumentIR(blocks=[List(items=[ListItem(text=f"项 {ok}")])])
        # One equation OLE, one decoded span (nested) → no retry needed.
        assert _mtef_recovery_needed(in_table, 1) is False
        assert _mtef_recovery_needed(in_list, 1) is False

    def test_auto_mode_no_retry_for_decoded_table_equation(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # End-to-end: a MathType OLE the native MTEF adapter decodes fine,
        # sitting in a table cell, must NOT drag "auto" into a LibreOffice
        # round-trip just because the span lives one level down.
        def fake_run(*args, **kwargs):
            raise AssertionError("LibreOffice should not be invoked")

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: "soffice",
        )
        monkeypatch.setattr("brailix.input.docx.subprocess.run", fake_run)

        path, doc = _make_docx(tmp_path)
        table = doc.add_table(rows=1, cols=1)
        para = table.cell(0, 0).paragraphs[0]
        para.add_run("公式 ")
        _embed_ole_equation(
            doc, para, _eqnolehdr_wrapped(_mtef_sample_payload())
        )
        doc.save(path)

        result = parse_docx(path, mathtype_fallback="auto", profile="cn_current", language="zh-CN")
        tables = [b for b in result.blocks if isinstance(b, Table)]
        assert tables
        cell_text = tables[0].rows[0].cells[0].text or ""
        assert "<msup>" in cell_text

    def test_auto_mode_retries_when_all_mtef_failed(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # When every OLE blob produces merror, "auto" should re-parse
        # via LibreOffice.
        calls: list[list[str]] = []

        def fake_run(cmd, *, check, capture_output, timeout):
            calls.append(cmd)
            out_dir = Path(cmd[cmd.index("--outdir") + 1])
            stem = Path(cmd[-1]).stem
            # Simulate LibreOffice replacing the OLE equation with a
            # native OMML one by rewriting the file to have a simple
            # OMML expression and no OLE.
            new = _make_docx(out_dir, name=f"{stem}.docx")
            _, conv = new
            para = conv.add_paragraph("修复后 ")
            para._p.append(_omml_fragment(
                '<m:sSup>'
                '<m:e><m:r><m:t>y</m:t></m:r></m:e>'
                '<m:sup><m:r><m:t>3</m:t></m:r></m:sup>'
                '</m:sSup>'
            ))
            conv.save(out_dir / f"{stem}.docx")

            class Result:
                returncode = 0
            return Result()

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: "soffice",
        )
        monkeypatch.setattr("brailix.input.docx.subprocess.run", fake_run)

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("前 ")
        # Garbled MTEF — adapter will emit merror.
        _embed_ole_equation(doc, para, bytes([5, 1, 0, 11, 0, 0xFF, 0xFF]))
        doc.save(path)

        result = parse_docx(path, mathtype_fallback="auto", profile="cn_current", language="zh-CN")
        assert len(calls) == 1
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "修复后" in joined
        # The fallback path swapped the bad OLE for a valid OMML, now a
        # deferred island that the frontend converts to ``<msup>``.
        islands = _inline_math_islands(joined)
        assert islands and "<msup>" in _island_mathml(islands[-1])
        assert "merror" not in joined

    def test_auto_mode_retries_when_equation_ole_vanishes_silently(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # An equation OLE whose blob is neither a CFB container nor raw
        # MTEF produces NO span at all (silent drop — corrupt container,
        # third-party writer, or olefile missing).  "auto" must treat
        # that as a failure to recover, not as "no equations here" —
        # the old heuristic ("no inline math → skip") lost the whole
        # formula silently.
        calls: list[list[str]] = []

        def fake_run(cmd, *, check, capture_output, timeout):
            calls.append(cmd)
            out_dir = Path(cmd[cmd.index("--outdir") + 1])
            stem = Path(cmd[-1]).stem
            new = _make_docx(out_dir, name=f"{stem}.docx")
            _, conv = new
            para = conv.add_paragraph("修复后 ")
            para._p.append(_omml_fragment(
                '<m:sSup>'
                '<m:e><m:r><m:t>y</m:t></m:r></m:e>'
                '<m:sup><m:r><m:t>3</m:t></m:r></m:sup>'
                '</m:sSup>'
            ))
            conv.save(out_dir / f"{stem}.docx")

            class Result:
                returncode = 0
            return Result()

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: "soffice",
        )
        monkeypatch.setattr("brailix.input.docx.subprocess.run", fake_run)

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("前 ")
        # First byte 0x00: not an EQNOLEFILEHDR, not a known MTEF
        # prelude, not a CFB container → _extract_mtef_payload returns
        # None and the walker emits nothing for this object.
        _embed_ole_equation(doc, para, b"\x00not-mtef-not-cfb")
        doc.save(path)

        result = parse_docx(path, mathtype_fallback="auto", profile="cn_current", language="zh-CN")
        assert len(calls) == 1
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "修复后" in joined

    def test_auto_mode_ignores_non_equation_ole(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # A non-formula OLE (chart / Excel sheet) must NOT trigger the
        # LibreOffice round-trip — there is no equation to recover, and
        # its blob never decodes, so without the ProgID gate this would
        # look exactly like the silent-loss case above.
        def fake_run(*args, **kwargs):
            raise AssertionError("LibreOffice should not be invoked")

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: "soffice",
        )
        monkeypatch.setattr("brailix.input.docx.subprocess.run", fake_run)

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("图表 ")
        _embed_ole_equation(
            doc, para, b"\x00chart-bytes", progid="Excel.Sheet.12"
        )
        doc.save(path)

        result = parse_docx(path, mathtype_fallback="auto", profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "图表" in joined

    def test_inner_dollar_is_escaped_in_inline_math(
        self, tmp_path: Path
    ) -> None:
        # A literal ``$`` inside a Word formula must not break the
        # ``$...$`` inline wrapping — the span used to terminate at the
        # inner dollar, corrupting the formula and leaking raw XML
        # fragments into the prose.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("前 ")
        para._p.append(_omml_fragment("<m:r><m:t>a$b</m:t></m:r>"))
        doc.save(path)

        result = parse_docx(path, profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        # The island's two wrappers are the only raw dollars; the inner one
        # is escaped inside the tagged payload, so it can't terminate the
        # ``$...$`` span early.
        assert joined.count("$") == 2
        islands = _inline_math_islands(joined)
        assert len(islands) == 1 and inline_math.is_tagged(islands[0])
        # The literal ``$`` round-trips: ``unwrap`` restores it in the raw
        # payload, and it survives into the converted MathML's text.
        import xml.etree.ElementTree as ET

        source, payload = inline_math.unwrap(islands[0])
        assert source == "omml" and "a$b" in payload
        tree = ET.fromstring(_island_mathml(islands[0]))
        assert "$" in "".join(tree.itertext())

    def test_auto_mode_swallows_libreoffice_unavailable(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # If LibreOffice isn't around, "auto" keeps the native result
        # (including merror) rather than raising — that's the point
        # of "auto" vs "libreoffice".
        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: None,
        )

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("前 ")
        _embed_ole_equation(doc, para, bytes([5, 1, 0, 11, 0, 0xFF, 0xFF]))
        doc.save(path)

        result = parse_docx(path, mathtype_fallback="auto", profile="cn_current", language="zh-CN")
        paragraphs = [b for b in result.blocks if isinstance(b, Paragraph)]
        joined = "\n".join(p.text or "" for p in paragraphs)
        assert "merror" in joined  # native fallback kept


# ---------------------------------------------------------------------------
# parse_doc legacy path
# ---------------------------------------------------------------------------


class TestParseDocLegacy:
    def test_missing_converter_raises_helpful_error(
        self, tmp_path: Path, monkeypatch
    ) -> None:
        # When no LibreOffice is on PATH, parse_doc must raise
        # ParseError with a "convert to .docx first" hint — not
        # silently fail or attempt to read the binary blob.
        from brailix.core.errors import ParseError

        monkeypatch.setattr(
            "brailix.input.docx._resolve_doc_converter",
            lambda override: None,
        )
        # Create a dummy .doc file so the FileNotFoundError path is
        # not hit before the converter check.
        path = tmp_path / "legacy.doc"
        path.write_bytes(b"\xd0\xcf\x11\xe0")  # OLE magic bytes
        with pytest.raises(ParseError) as exc:
            parse_doc(path, profile="cn_current", language="zh-CN")
        assert ".doc" in str(exc.value)
        assert "LibreOffice" in str(exc.value) or "soffice" in str(exc.value)


# ---------------------------------------------------------------------------
# parse_file suffix dispatch
# ---------------------------------------------------------------------------


class TestParseFileDispatch:
    def test_docx_routes_through_parse_docx(self, tmp_path: Path) -> None:
        # End-to-end via parse_file: ``.docx`` suffix dispatches to the
        # docx adapter (not the plain reader, which would choke on
        # binary content).
        from brailix.input import parse_file

        path, doc = _make_docx(tmp_path)
        doc.add_heading("章节", level=1)
        doc.add_paragraph("正文。")
        doc.save(path)

        result = parse_file(path, profile="cn_current", language="zh-CN")
        headings = [b for b in result.blocks if isinstance(b, Heading)]
        assert len(headings) == 1
        assert headings[0].text == "章节"

    def test_docm_suffix_also_routes_to_docx(self, tmp_path: Path) -> None:
        # ``.docm`` (macro-enabled) is the same OOXML container with
        # macros embedded; parse_docx reads it identically.
        from brailix.input import parse_file

        path = tmp_path / "macros.docm"
        doc = Document()
        doc.add_paragraph("hi")
        doc.save(path)

        result = parse_file(path, profile="cn_current", language="zh-CN")
        # Just confirm the parser ran (no exception).
        assert result.blocks  # not empty


# ---------------------------------------------------------------------------
# Error paths
# ---------------------------------------------------------------------------


class TestErrorPaths:
    def test_missing_file_raises(self, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            parse_docx(tmp_path / "nope.docx", profile="cn_current", language="zh-CN")

    def test_not_a_zip_raises_parse_error(self, tmp_path: Path) -> None:
        from brailix.core.errors import ParseError

        bogus = tmp_path / "fake.docx"
        bogus.write_bytes(b"not a real zip archive")
        with pytest.raises(ParseError):
            parse_docx(bogus, profile="cn_current", language="zh-CN")


# ---------------------------------------------------------------------------
# Run-level vertical alignment (Ctrl+Shift+= / Ctrl+= scripts) → inline math
# ---------------------------------------------------------------------------


def _vert_run_xml(text: str, vert: str | None) -> str:
    """One ``<w:r>`` with an optional ``<w:vertAlign>`` (the property Word
    sets for super/subscript)."""
    rpr = f'<w:rPr><w:vertAlign w:val="{vert}"/></w:rPr>' if vert else ""
    return (
        f'<w:r xmlns:w="{_W_NS}">{rpr}'
        f'<w:t xml:space="preserve">{text}</w:t></w:r>'
    )


def _append_script_runs(paragraph, runs: list[tuple[str, str | None]]) -> None:
    """Append ``(text, vert)`` runs to ``paragraph``; ``vert`` is
    ``"superscript"`` / ``"subscript"`` / ``None``."""
    runs_xml = "".join(_vert_run_xml(t, v) for t, v in runs)
    wrapper = etree.fromstring(f'<root xmlns:w="{_W_NS}">{runs_xml}</root>')
    for run in list(wrapper):
        paragraph._p.append(run)


def _para_text(result) -> str:
    return "\n".join(
        p.text or "" for p in result.blocks if isinstance(p, Paragraph)
    )


class TestVertAlignScripts:
    """Scripts set via Ctrl+Shift+= / Ctrl+= become inline math. Always on —
    they used to be flattened to bare text, losing the formula."""

    def test_superscript_becomes_msup(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(para, [("x", None), ("2", "superscript")])
        doc.save(path)
        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        # Deferred: the cluster is a linearised ``script_cluster`` island the
        # frontend converts to ``<msup>`` — input builds no MathML itself.
        mathml = _island_mathml(_inline_math_islands(text)[0])
        assert "<msup>" in mathml
        assert "<mi>x</mi>" in mathml and "<mn>2</mn>" in mathml

    def test_subscript_becomes_msub(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(
            para, [("H", None), ("2", "subscript"), ("O", None)]
        )
        doc.save(path)
        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        mathml = _island_mathml(_inline_math_islands(text)[0])
        assert "<msub>" in mathml
        assert "<mi>H</mi>" in mathml and "<mn>2</mn>" in mathml
        assert "<mi>O</mi>" in mathml
        # chemistry detection is off by default → plain math, no chem tag.
        assert "data-bk-chem" not in mathml

    def test_prose_around_script_stays_prose(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        para.add_run("面积是 ")
        _append_script_runs(para, [("x", None), ("2", "superscript")])
        para.add_run(" 平方米")
        doc.save(path)
        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "面积是" in text and "平方米" in text
        islands = _inline_math_islands(text)
        assert len(islands) == 1
        # The surrounding Chinese must not be swallowed into the math island.
        assert "面积" not in islands[0] and "平方" not in islands[0]
        assert "<msup>" in _island_mathml(islands[0])

    def test_unit_superscript(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)  # m² — square metre
        para = doc.add_paragraph()
        _append_script_runs(para, [("m", None), ("2", "superscript")])
        doc.save(path)
        mathml = _island_mathml(_inline_math_islands(_para_text(parse_docx(path, profile="cn_current", language="zh-CN")))[0])
        assert "<msup>" in mathml and "<mi>m</mi>" in mathml

    def test_subscript_then_superscript_is_msubsup(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(
            para, [("x", None), ("1", "subscript"), ("2", "superscript")]
        )
        doc.save(path)
        islands = _inline_math_islands(_para_text(parse_docx(path, profile="cn_current", language="zh-CN")))
        assert "<msubsup>" in _island_mathml(islands[0])

    def test_negative_exponent_uses_canonical_minus(self, tmp_path: Path) -> None:
        # 10⁻³ — the hyphen-minus is canonicalised to U+2212 so the math
        # backend's symbol table matches (a raw '-' is MATH_UNKNOWN_SYMBOL).
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(
            para, [("10", None), ("-", "superscript"), ("3", "superscript")]
        )
        doc.save(path)
        mathml = _island_mathml(_inline_math_islands(_para_text(parse_docx(path, profile="cn_current", language="zh-CN")))[0])
        assert "<msup>" in mathml and "−" in mathml and "<mn>3</mn>" in mathml

    def test_plain_text_without_vertalign_unchanged(self, tmp_path: Path) -> None:
        # Regression: ordinary "H2O" (no vertAlign) must NOT become math.
        path, doc = _make_docx(tmp_path)
        doc.add_paragraph("H2O 是水")
        doc.save(path)
        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "$<math" not in text and "H2O" in text

    def test_super_run_without_text_is_safe(self, tmp_path: Path) -> None:
        # A superscript run carrying no <w:t> (e.g. an auto footnote mark)
        # produces no math and doesn't crash.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        para.add_run("见正文")
        empty = (
            f'<w:r xmlns:w="{_W_NS}"><w:rPr>'
            f'<w:vertAlign w:val="superscript"/></w:rPr></w:r>'
        )
        wrapper = etree.fromstring(f'<root xmlns:w="{_W_NS}">{empty}</root>')
        for run in list(wrapper):
            para._p.append(run)
        doc.save(path)
        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "见正文" in text and "$<math" not in text

    def test_cluster_adjacent_to_omml_no_double_dollar(self, tmp_path: Path) -> None:
        # A script cluster directly followed by an OMML island must keep a
        # separator so the segmenter's inline-math regex (which rejects $$)
        # still sees both.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(para, [("a", None), ("2", "superscript")])
        para._p.append(_omml_fragment("<m:r><m:t>z</m:t></m:r>"))
        doc.save(path)
        assert "$$" not in _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))

    def test_end_to_end_braille_superscript(self, tmp_path: Path) -> None:
        from brailix.pipeline import Pipeline

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(para, [("x", None), ("2", "superscript")])
        doc.save(path)
        result = Pipeline(profile="cn_current").translate_document(parse_docx(path, profile="cn_current", language="zh-CN"))
        # ⠰⠭⠌⠆ : latin-x prefix, superscript marker, lowered 2.
        assert "⠰⠭⠌⠆" in result.render("unicode")


class TestVertAlignChemistry:
    """Conservative chemistry reading of script clusters — opt-in via
    ``chem_detection`` (the ``input.docx.detect_chemistry`` profile feature),
    off by default because a lone single-letter subscript coincides with an
    element symbol."""

    def test_off_by_default_h2o_is_math(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(
            para, [("H", None), ("2", "subscript"), ("O", None)]
        )
        doc.save(path)
        mathml = _island_mathml(_inline_math_islands(_para_text(parse_docx(path, profile="cn_current", language="zh-CN")))[0])
        assert "data-bk-chem" not in mathml and "<msub>" in mathml

    def test_on_h2o_tagged_chem(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(
            para, [("H", None), ("2", "subscript"), ("O", None)]
        )
        doc.save(path)
        text = _para_text(parse_docx(path, chem_detection=True, profile="cn_current", language="zh-CN"))
        assert 'data-bk-chem="1"' in _island_mathml(_inline_math_islands(text)[0])

    def test_single_letter_variable_stays_math(self, tmp_path: Path) -> None:
        # V₁ — vanadium IS an element, but a lone single-letter subscript is
        # almost always a maths/physics variable, so it stays math.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(para, [("V", None), ("1", "subscript")])
        doc.save(path)
        mathml = _island_mathml(
            _inline_math_islands(_para_text(parse_docx(path, chem_detection=True, profile="cn_current", language="zh-CN")))[0]
        )
        assert "data-bk-chem" not in mathml and "<msub>" in mathml

    def test_lowercase_base_stays_math(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)  # x² — x isn't an element symbol
        para = doc.add_paragraph()
        _append_script_runs(para, [("x", None), ("2", "superscript")])
        doc.save(path)
        mathml = _island_mathml(
            _inline_math_islands(_para_text(parse_docx(path, chem_detection=True, profile="cn_current", language="zh-CN")))[0]
        )
        assert "data-bk-chem" not in mathml and "<msup>" in mathml

    def test_charge_is_chem(self, tmp_path: Path) -> None:
        path, doc = _make_docx(tmp_path)  # Ca²⁺
        para = doc.add_paragraph()
        _append_script_runs(
            para,
            [("Ca", None), ("2", "superscript"), ("+", "superscript")],
        )
        doc.save(path)
        text = _para_text(parse_docx(path, chem_detection=True, profile="cn_current", language="zh-CN"))
        assert 'data-bk-chem="1"' in _island_mathml(_inline_math_islands(text)[0])

    def test_end_to_end_chem_braille_h2o(self, tmp_path: Path) -> None:
        from brailix.pipeline import Pipeline

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(
            para, [("H", None), ("2", "subscript"), ("O", None)]
        )
        doc.save(path)
        doc_ir = parse_docx(path, chem_detection=True, profile="cn_current", language="zh-CN")
        result = Pipeline(profile="cn_current").translate_document(doc_ir)
        # ⠸⠓⠆⠕ : chemical-formula indicator ⠸, H, lowered 2, O.
        assert "⠸⠓⠆⠕" in result.render("unicode")

    def test_profile_feature_enables_chem(self, tmp_path: Path) -> None:
        # cn_current ships input.docx.detect_chemistry=true, so the Pipeline
        # reads H₂O as chemistry through translate_file without the caller
        # passing chem_detection at all.
        from brailix.pipeline import Pipeline

        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        _append_script_runs(
            para, [("H", None), ("2", "subscript"), ("O", None)]
        )
        doc.save(path)
        result = Pipeline(profile="cn_current").translate_file(path)
        assert "⠸⠓⠆⠕" in result.render("unicode")


# ---------------------------------------------------------------------------
# w:jc relative spellings (Word 2013+ writes start / end, not left / right)
# ---------------------------------------------------------------------------


class TestJcAlignmentValues:
    """``_paragraph_alignment`` maps the LTR-relative ``w:jc`` spellings Word
    2013+ emits. python-docx's alignment enum only round-trips
    left / right / center / both, so the ``start`` / ``end`` branch needs raw
    XML to reach."""

    def _para(self, jc_val: str) -> etree._Element:
        return etree.fromstring(
            f'<w:p xmlns:w="{_W_NS}"><w:pPr>'
            f'<w:jc w:val="{jc_val}"/></w:pPr></w:p>'
        )

    def test_end_maps_to_right(self) -> None:
        from brailix.input.docx._blocks import _paragraph_alignment

        assert _paragraph_alignment(self._para("end")) == "right"

    def test_start_has_no_align(self) -> None:
        # ``start`` is the LTR default (flush-left) — no braille marker.
        from brailix.input.docx._blocks import _paragraph_alignment

        assert _paragraph_alignment(self._para("start")) is None

    def test_bare_val_without_prefix_is_read(self) -> None:
        # Some emitters drop the ``w:`` prefix; the value must still resolve
        # (exercises the _ns_attr bare-name fallback on w:jc).
        from brailix.input.docx._blocks import _paragraph_alignment

        p = etree.fromstring(
            f'<w:p xmlns:w="{_W_NS}"><w:pPr><w:jc val="center"/></w:pPr></w:p>'
        )
        assert _paragraph_alignment(p) == "center"


# ---------------------------------------------------------------------------
# Run-level <w:br> / <w:tab> and hyperlink-wrapped content
# ---------------------------------------------------------------------------


class TestRunBreaksAndHyperlink:
    def test_break_and_tab_in_run(self, tmp_path: Path) -> None:
        # <w:br> contributes a newline, <w:tab> a space, inside paragraph text.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        runs_xml = (
            f'<w:r xmlns:w="{_W_NS}"><w:t>甲</w:t><w:br/><w:t>乙</w:t></w:r>'
            f'<w:r xmlns:w="{_W_NS}"><w:t>丙</w:t><w:tab/><w:t>丁</w:t></w:r>'
        )
        wrapper = etree.fromstring(f'<root xmlns:w="{_W_NS}">{runs_xml}</root>')
        for r in list(wrapper):
            para._p.append(r)
        doc.save(path)

        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "甲\n乙" in text   # <w:br> → newline
        assert "丙 丁" in text     # <w:tab> → space

    def test_hyperlink_wrapped_run_text_surfaces(self, tmp_path: Path) -> None:
        # A <w:hyperlink> wraps runs; their text must reach the paragraph.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph("看 ")
        hyper_xml = (
            f'<w:hyperlink xmlns:w="{_W_NS}"><w:r><w:t>链接文字</w:t></w:r>'
            f'</w:hyperlink>'
        )
        para._p.append(etree.fromstring(hyper_xml))
        doc.save(path)

        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        assert "看" in text and "链接文字" in text

    def test_hyperlink_wrapped_math_surfaces(self, tmp_path: Path) -> None:
        # Inline OMML inside a hyperlink is still extracted as inline math.
        path, doc = _make_docx(tmp_path)
        para = doc.add_paragraph()
        hyper_xml = (
            f'<w:hyperlink xmlns:w="{_W_NS}" xmlns:m="{_M_NS}">'
            f'<m:oMath><m:sSup>'
            f'<m:e><m:r><m:t>x</m:t></m:r></m:e>'
            f'<m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
            f'</m:sSup></m:oMath>'
            f'</w:hyperlink>'
        )
        para._p.append(etree.fromstring(hyper_xml))
        doc.save(path)

        text = _para_text(parse_docx(path, profile="cn_current", language="zh-CN"))
        # Inline OMML inside a hyperlink defers like any other inline OMML.
        islands = _inline_math_islands(text)
        assert len(islands) == 1 and inline_math.unwrap(islands[0])[0] == "omml"
        assert "<msup>" in _island_mathml(islands[0])
